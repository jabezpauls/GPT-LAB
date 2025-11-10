import gradio as gr
import requests
from PyPDF2 import PdfReader
from docx import Document
import json
import os
import time

OLLAMA_URL = "http://localhost:11434/api/generate"
MODEL_NAME = "qwen2.5:0.5b"

# Text Extraction
def extract_text(file):
    """Extract text from PDF, DOCX, or TXT files."""
    if file is None:
        return None

    filename = file.name
    text = ""
    try:
        if filename.endswith(".pdf"):
            reader = PdfReader(filename)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        elif filename.endswith(".docx"):
            doc = Document(filename)
            text = "\n".join([p.text for p in doc.paragraphs if p.text.strip() != ""])
        elif filename.endswith(".txt"):
            with open(filename, "r", encoding="utf-8") as f:
                text = f.read()
        else:
            return None
        return text.strip()
    except Exception as e:
        return None

# Ollama Helper
def ollama_generate(prompt, max_tokens=1000):
    """Generate text using Ollama API."""
    try:
        response = requests.post(
            OLLAMA_URL,
            json={
                "model": MODEL_NAME,
                "prompt": prompt,
                "stream": True,
                "temperature": 0.2,
                "max_tokens": max_tokens
            },
            stream=True
        )
        if response.status_code != 200:
            return f"Error: {response.status_code} - {response.text}"

        full_response = ""
        for line in response.iter_lines():
            if line:
                try:
                    data = json.loads(line.decode("utf-8"))
                    if "response" in data:
                        full_response += data["response"]
                except json.JSONDecodeError:
                    continue
        return full_response.strip() if full_response else "No response content."
    except Exception as e:
        return f"Error connecting to Ollama API: {e}"

# Summarization Prompt Builder
def build_summary_prompt(text, summary_level, doc_type):
    """Build prompt based on summary level and document type."""
    if summary_level == "Short":
        return f"Summarize the following {doc_type} in 100 words (no preamble, no markdown, output only summary):\n\n{text}"
    elif summary_level == "Medium":
        return f"Summarize the following {doc_type} in 200 words highlighting key points (no preamble no markdown, output only summary):\n\n{text}"
    else:  # Detailed
        if doc_type == "Legal":
            return f"Provide a detailed section-wise summary of this legal document. Include Parties, Payment Terms, Liabilities, Confidentiality, Risks, Termination (no preamble no markdown, output only summary):\n\n{text}"
        else:
            return f"Provide a detailed section-wise summary of this literary text. Include Characters, Plot, Themes, Moral Lessons (no preamble no markdown, output only summary):\n\n{text}"

# Key Clause / Theme Extraction
def extract_keywords(text, doc_type):
    """Extract keywords, clauses, or themes from document."""
    if doc_type == "Legal":
        prompt = f"Extract key clauses from this legal document. Include Payment Terms, Liabilities, Confidentiality, Risks, Termination.\n\n{text}"
    else:
        prompt = f"Extract key elements from this literary text. Include main Characters, Plot Events, Themes, Moral Lessons.\n\n{text}"
    return ollama_generate(prompt, max_tokens=400)

# Q&A Prompt Builder
def answer_question(text, user_question):
    """Answer user questions based on document content."""
    prompt = f"""
Answer the following question accurately based on the document content below:
Content:
{text}
Question: {user_question}
"""
    return ollama_generate(prompt, max_tokens=500)

# Export Function
def export_results(summary, qa, keywords, export_format="txt"):
    """Export analysis results to file."""
    filename = f"document_analysis.{export_format}"
    try:
        if export_format == "txt":
            with open(filename, "w", encoding="utf-8") as f:
                f.write("---Summary---\n" + summary + "\n\n")
                f.write("---Q&A---\n" + qa + "\n\n")
                f.write("---Keywords / Clauses---\n" + keywords)
        elif export_format == "docx":
            doc = Document()
            doc.add_heading("AI Document Analysis Report", 0)
            doc.add_paragraph("---Summary---\n" + summary)
            doc.add_paragraph("---Q&A---\n" + qa)
            doc.add_paragraph("---Keywords / Clauses---\n" + keywords)
            doc.save(filename)
        elif export_format == "pdf":
            from reportlab.lib.pagesizes import A4
            from reportlab.pdfgen import canvas
            c = canvas.Canvas(filename, pagesize=A4)
            c.setFont("Helvetica", 10)
            y = 800
            for section in ["Summary", "Q&A", "Keywords / Clauses"]:
                c.drawString(50, y, f"---{section}---")
                y -= 20
                content = {"Summary": summary, "Q&A": qa, "Keywords / Clauses": keywords}[section]
                for line in content.split("\n"):
                    c.drawString(50, y, line[:90])
                    y -= 15
                    if y < 50:
                        c.showPage()
                        y = 800
            c.save()
        return filename
    except Exception as e:
        return f"Export failed: {str(e)}"

# Main Processing Function
def analyze_document(file, doc_type, summary_level, user_question, export_format, progress=gr.Progress()):
    """Main analysis function with progress tracking."""
    if file is None:
        return "❌ Please upload a document first.", "", "", "", "0", "0s"

    progress(0.1, desc="Extracting text from document...")
    start_time = time.time()

    text = extract_text(file)
    if text is None:
        return "❌ Error: Unable to extract text. Ensure file is not corrupted and format is supported.", "", "", "", "0", "0s"

    word_count = len(text.split())
    char_count = len(text)

    progress(0.3, desc="Generating summary...")
    summary = ollama_generate(build_summary_prompt(text, summary_level, doc_type))

    progress(0.5, desc="Extracting keywords/clauses...")
    keywords = extract_keywords(text, doc_type)

    qa = ""
    if user_question and user_question.strip() != "":
        progress(0.7, desc="Answering your question...")
        qa = answer_question(text, user_question)

    progress(0.9, desc="Exporting results...")
    exported_file = export_results(summary, qa if qa else "No question asked", keywords, export_format)

    end_time = time.time()
    processing_time = f"{end_time - start_time:.2f}s"

    progress(1.0, desc="Complete!")

    export_message = f"✅ Analysis complete! File exported: {exported_file}"

    return summary, qa, keywords, export_message, str(word_count), processing_time

# Custom CSS for clean, modern design
custom_css = """
.gradio-container {
    font-family: 'Inter', -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif;
    max-width: 1400px;
    margin: auto;
}

.header {
    text-align: center;
    padding: 2.5rem 0;
    background: linear-gradient(135deg, #0f172a 0%, #1e293b 100%);
    color: white;
    border-radius: 12px;
    margin-bottom: 2rem;
    box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1);
}

.stats-card {
    background: #f8fafc;
    border-radius: 8px;
    padding: 1rem;
    border: 1px solid #e2e8f0;
    text-align: center;
}

.analyze-btn {
    background: linear-gradient(135deg, #0f172a 0%, #334155 100%) !important;
    color: white !important;
    font-weight: 600 !important;
    padding: 0.75rem 2rem !important;
    border-radius: 8px !important;
    transition: all 0.3s ease !important;
}

.analyze-btn:hover {
    transform: translateY(-2px) !important;
    box-shadow: 0 4px 12px rgba(0, 0, 0, 0.15) !important;
}

.output-section {
    background: white;
    border-radius: 8px;
    padding: 1.5rem;
    border: 1px solid #e2e8f0;
    box-shadow: 0 1px 3px rgba(0, 0, 0, 0.05);
}
"""

# Gradio Interface
with gr.Blocks(css=custom_css, theme=gr.themes.Soft(), title="AI Document Analyzer") as demo:

    # Header
    gr.HTML("""
        <div class="header">
            <h1 style="margin:0; font-size: 2.5rem;">📄 AI Document Analyzer</h1>
            <p style="margin:0.5rem 0 0 0; opacity: 0.9;">Summarize, Query, and Export Legal & Literary Documents</p>
        </div>
    """)

    with gr.Row():
        with gr.Column(scale=1):
            gr.Markdown("### 📤 Upload & Configuration")

            file_input = gr.File(
                file_types=[".pdf", ".docx", ".txt"],
                label="Upload Document",
                file_count="single"
            )

            with gr.Row():
                doc_type_input = gr.Radio(
                    ["Legal", "Literary"],
                    value="Legal",
                    label="📋 Document Type",
                    info="Select document category"
                )

                summary_level_input = gr.Radio(
                    ["Short", "Medium", "Detailed"],
                    value="Medium",
                    label="📊 Summary Level",
                    info="Choose summary depth"
                )

            export_format_input = gr.Radio(
                ["txt", "docx", "pdf"],
                value="txt",
                label="💾 Export Format",
                info="Select output format"
            )

            user_question_input = gr.Textbox(
                lines=3,
                placeholder="e.g., What are the payment terms? Who are the main characters?",
                label="❓ Ask a Question (Optional)",
                info="Query specific information from the document"
            )

            analyze_button = gr.Button("🚀 Analyze Document", elem_classes=["analyze-btn"], size="lg")

            gr.Markdown("---")

            # Statistics
            gr.Markdown("### 📊 Document Statistics")
            with gr.Row():
                words_stat = gr.Textbox(label="📝 Words", interactive=False, scale=1)
                time_stat = gr.Textbox(label="⏱️ Time", interactive=False, scale=1)

        with gr.Column(scale=2):
            gr.Markdown("### 📑 Analysis Results")

            with gr.Tabs():
                with gr.TabItem("📄 Summary"):
                    summary_output = gr.Textbox(
                        lines=18,
                        label="",
                        show_label=False,
                        placeholder="Document summary will appear here...",
                        interactive=False,
                        elem_classes=["output-section"]
                    )

                with gr.TabItem("💬 Q&A"):
                    qa_output = gr.Textbox(
                        lines=18,
                        label="",
                        show_label=False,
                        placeholder="Question answers will appear here...",
                        interactive=False,
                        elem_classes=["output-section"]
                    )

                with gr.TabItem("🔑 Keywords / Key Clauses"):
                    keywords_output = gr.Textbox(
                        lines=18,
                        label="",
                        show_label=False,
                        placeholder="Extracted keywords and clauses will appear here...",
                        interactive=False,
                        elem_classes=["output-section"]
                    )

                with gr.TabItem("💾 Export Status"):
                    export_output = gr.Textbox(
                        lines=18,
                        label="",
                        show_label=False,
                        placeholder="Export status will appear here...",
                        interactive=False,
                        elem_classes=["output-section"]
                    )

    # Footer
    gr.Markdown("""
    ---
    <div style="text-align: center; color: #64748b; font-size: 0.9rem;">
        <p>💡 <strong>Supported Formats:</strong> PDF, DOCX, TXT | <strong>Model:</strong> Qwen 2.5 (0.5b) via Ollama</p>
        <p>📌 <strong>Tip:</strong> Upload clear, text-based documents for best results</p>
    </div>
    """)

    # Connect button to function
    analyze_button.click(
        fn=analyze_document,
        inputs=[file_input, doc_type_input, summary_level_input, user_question_input, export_format_input],
        outputs=[summary_output, qa_output, keywords_output, export_output, words_stat, time_stat]
    )


if __name__ == "__main__":
    print("=" * 60)
    print("🚀 Starting AI Document Analyzer")
    print("=" * 60)
    print(f"📋 Model: {MODEL_NAME}")
    print(f"🌐 Ollama URL: {OLLAMA_URL}")
    print("📄 Supported formats: PDF, DOCX, TXT")
    print("=" * 60)
    print("\n⚠️  Ensure Ollama is running with: ollama pull qwen2.5:0.5b\n")
    demo.launch(share=False, show_error=True)
